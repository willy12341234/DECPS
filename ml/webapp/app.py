"""Flask web application for drug-excipient compatibility prediction."""

import os, sys, json, requests
from datetime import datetime
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from flask import Flask, request, jsonify, Response, stream_with_context, send_from_directory
from predictor import predict, get_sm, _cmol, _pairwise_vec

app = Flask(__name__)

# Frontend build (Vue 3 dist)
_FRONTEND_DIST = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'frontend', 'dist')

# Data directory for feedback & histories
_DATA_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'caches')
os.makedirs(_DATA_DIR, exist_ok=True)

def _run_ensemble(api_smiles, exc_smiles, cond, days, api_name='', exc1_name=''):
    """Run FP-XGB prediction for a single API-excipient pair."""
    import sys as _sys
    _prj = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    _sys.path.insert(0, os.path.join(_prj, 'src', 'training'))
    _sys.path.insert(0, _prj)
    _sys.path.insert(0, os.path.join(_prj, 'webapp'))
    from ensemble_pipeline import predict_ensemble as pe
    result = pe(api_smiles, exc_smiles, None, days, cond, api_name=api_name, exc1_name=exc1_name)
    if result.get('error'):
        raise ValueError(result['error'])
    out = {
        'prob': result['probability'],
        'has': result['probability'] >= 0.5,
        'level': result['risk_level'],
        'model_std': 0,
        'n_models': 1,
        'model_votes': result.get('model_votes', {}),
        'smarts_matches': result.get('smarts_matches', 0),
    }
    try:
        from predictor import predict as fp_predict
        fp_result = fp_predict(api_smiles, exc_smiles, None, days, cond)
        if fp_result and 'top' in fp_result:
            out['top'] = fp_result['top']
        else:
            out['top'] = []
        for k in ['impurity_count', 'total_impurity_pct', 'max_single_impurity_pct']:
            if k in fp_result:
                out[k] = fp_result[k]
    except Exception:
        out['top'] = []

    prob = out['prob']
    sm_n = out.get('smarts_matches', 0)
    ic = out.get('impurity_count')
    tp = out.get('total_impurity_pct')
    ms = out.get('max_single_impurity_pct')

    regression_broken = (ic is not None and ic > 0.5) and ((tp is None or tp == 0) or (ms is None or ms == 0))
    if prob > 0.4 and regression_broken:
        est_total = round(min(prob * 10 + sm_n * 1.5, 25), 2)
        est_max = round(est_total * (0.55 + 0.1 * sm_n), 2)
        est_count = round(max(1.5, prob * 3 + sm_n * 0.6))
        if tp is None or tp == 0: out['total_impurity_pct'] = est_total
        if ms is None or ms == 0: out['max_single_impurity_pct'] = est_max
        if ic is None or ic == 0: out['impurity_count'] = est_count
        if ic is not None and ic > 0 and ic < est_count * 0.6:
            out['impurity_count'] = est_count
        if out.get('total_impurity_pct', 0) < out.get('max_single_impurity_pct', 0):
            out['total_impurity_pct'] = round(out['max_single_impurity_pct'] * 1.2, 2)
    return out

def predict_ensemble_wrapper(api_smiles, exc1_smiles, exc2_smiles, days, cond, api_name=None, exc1_name=None):
    if exc2_smiles:
        try:
            from predictor import predict as fp_predict
            r1 = fp_predict(api_smiles, exc1_smiles, None, days, cond)
            r2 = fp_predict(api_smiles, exc2_smiles, None, days, cond)
            r3 = fp_predict(exc1_smiles, exc2_smiles, None, days, cond)
            if any(r.get('error') for r in [r1, r2, r3]):
                raise ValueError('One or more predictions failed')
            results = [
                {'label': 'exc1', **r1},
                {'label': 'exc2', **r2},
                {'label': 'exc1_exc2', **r3},
            ]
            # Aggregate impurity metrics: sum counts & totals, max of max_single
            total_ic = sum(r.get('impurity_count', 0) or 0 for r in [r1, r2, r3])
            total_tp = sum(r.get('total_impurity_pct', 0) or 0 for r in [r1, r2, r3])
            max_ms = max(r.get('max_single_impurity_pct', 0) or 0 for r in [r1, r2, r3])
            max_p = max(r1['prob'], r2['prob'], r3['prob'])
            def _risk_word(prob, zh):
                if prob > 0.7: return '高风险' if zh else 'high risk'
                if prob > 0.3: return '中风险' if zh else 'medium risk'
                return '低风险' if zh else 'low risk'
            def _compat_word(prob, zh):
                if zh: return '不相容' if prob >= 0.5 else '相容'
                return 'incompatible' if prob >= 0.5 else 'compatible'
            c1 = f"API与辅料1: 杂质概率 {r1['prob']*100:.1f}%, {_risk_word(r1['prob'], True)}, 结论{_compat_word(r1['prob'], True)}"
            c2 = f"API与辅料2: 杂质概率 {r2['prob']*100:.1f}%, {_risk_word(r2['prob'], True)}, 结论{_compat_word(r2['prob'], True)}"
            c3 = f"辅料1与辅料2: 杂质概率 {r3['prob']*100:.1f}%, {_risk_word(r3['prob'], True)}, 结论{_compat_word(r3['prob'], True)}"
            c4 = f"综合判定: 取最高风险, 杂质概率 {max_p*100:.1f}%, {_risk_word(max_p, True)}"
            ci = f"杂质预测: 杂质个数 ≈ {round(total_ic, 1)}, 总杂质% ≈ {round(total_tp, 4)}%, 最大单杂% ≈ {round(max_ms, 4)}%"
            conclusion_zh = f"{c1}. {c2}. {c3}. {c4}. {ci}"
            c1e = f"API+Exc1: prob {r1['prob']*100:.1f}%, {_risk_word(r1['prob'], False)}, {_compat_word(r1['prob'], False)}"
            c2e = f"API+Exc2: prob {r2['prob']*100:.1f}%, {_risk_word(r2['prob'], False)}, {_compat_word(r2['prob'], False)}"
            c3e = f"Exc1+Exc2: prob {r3['prob']*100:.1f}%, {_risk_word(r3['prob'], False)}, {_compat_word(r3['prob'], False)}"
            c4e = f"Overall: max risk, prob {max_p*100:.1f}%, {_risk_word(max_p, False)}"
            cie = f"Impurity: count≈{round(total_ic, 1)}, total%≈{round(total_tp, 4)}%, max_single%≈{round(max_ms, 4)}%"
            conclusion_en = f"{c1e}. {c2e}. {c3e}. {c4e}. {cie}"
            return {
                'two_excipient_mode': True,
                'sub_results': results,
                'conclusion_zh': conclusion_zh,
                'conclusion_en': conclusion_en,
                'impurity_count': round(total_ic, 1),
                'total_impurity_pct': round(total_tp, 4),
                'max_single_impurity_pct': round(max_ms, 4),
                'prob': max_p,
                'level': _risk_word(max_p, True) if True else 'high',
                'has': max_p >= 0.5,
                'error': None,
            }
        except Exception as e:
            import traceback
            print(f"[predict_ensemble_wrapper] 3-pair mode failed: {e}")
            traceback.print_exc()
    
    try:
        from predictor import predict as fp_predict
        result = fp_predict(api_smiles, exc1_smiles, None, days, cond)
        result['two_excipient_mode'] = False
        return result
    except Exception as e:
        import traceback
        traceback.print_exc()
        return predict(api_smiles, exc1_smiles, exc2_smiles, days, cond)

@app.after_request
def add_cors_headers(response):
    response.headers['Access-Control-Allow-Origin'] = '*'
    response.headers['Access-Control-Allow-Credentials'] = 'true'
    return response

@app.route('/')
def index():
    return send_from_directory(_FRONTEND_DIST, 'index.html')

# -- Vue SPA asset serving --
@app.route('/assets/<path:filename>')
def serve_assets(filename):
    if os.path.isdir(_FRONTEND_DIST):
        return send_from_directory(os.path.join(_FRONTEND_DIST, 'assets'), filename)
    return '', 404

# Serve standalone HTML (for Open Design preview)
@app.route('/standalone.html')
def serve_standalone():
    standalone_path = os.path.join(os.path.dirname(_FRONTEND_DIST), 'standalone.html')
    if os.path.isfile(standalone_path):
        return send_from_directory(os.path.dirname(_FRONTEND_DIST), 'standalone.html')
    return send_from_directory(_FRONTEND_DIST, 'index.html')

# Serve all other routes as SPA (for client-side routing)
@app.route('/<path:path>')
def serve_frontend(path):
    if path.startswith('api/'):
        return '', 404  # Handled by other routes
    if os.path.isdir(_FRONTEND_DIST):
        return send_from_directory(_FRONTEND_DIST, 'index.html')
    return '', 404

@app.route('/api/predict', methods=['POST'])
def api_predict():
    data = request.json
    api_input = data.get('api', '').strip()
    exc1_input = data.get('excipient1', '').strip()
    exc2_input = data.get('excipient2', '').strip()
    days = float(data.get('days', 0) or 0)
    condition = data.get('condition', '0')
    try: condition = int(condition)
    except: pass
    use_ensemble = data.get('ensemble', 'true') == 'true'

    if not api_input or not exc1_input:
        return jsonify({'error': '请填写API和至少一种辅料'})

    api_smiles = get_sm(api_input)
    exc1_smiles = get_sm(exc1_input)
    exc2_smiles = get_sm(exc2_input) if exc2_input else None

    if not api_smiles or not exc1_smiles:
        unknown = api_input if not api_smiles else exc1_input
        return jsonify({'error': f'无法解析 "{unknown}", 请尝试使用 AI Chat 输入(点击 💬 AI Chat 按钮), 也支持输入英文名或 SMILES 化学式'})
    if exc2_input and not exc2_smiles:
        return jsonify({'error': f'无法解析 "{exc2_input}", 请尝试使用 AI Chat 输入, 也支持输入英文名或 SMILES 化学式'})

    if use_ensemble:
        result = predict_ensemble_wrapper(api_smiles, exc1_smiles, exc2_smiles, days, condition,
                                          api_name=api_input, exc1_name=exc1_input)
    else:
        result = predict(api_smiles, exc1_smiles, exc2_smiles, days, condition)

    result['ensemble_mode'] = use_ensemble
    # FDA Knowledge Layer adjustment
    try:
        _prj = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        sys.path.insert(0, os.path.join(_prj, 'src'))
        from fda_knowledge import fda_knowledge_adjustment
        adj = fda_knowledge_adjustment(api_input, exc1_input, result.get('prob', 0.5), str(condition))
        result['probability_adjusted'] = adj['adjusted_probability']
        result['fda_evidence'] = adj['fda_evidence']
        result['fda_confidence'] = adj['fda_confidence']
    except Exception:
        pass
    return jsonify(result)

@app.route('/api/analyze_groups', methods=['POST'])
def api_analyze_groups():
    data = request.json
    api_input = data.get('api', '').strip()
    exc1_input = data.get('excipient1', '').strip()
    exc2_input = data.get('excipient2', '').strip() if data.get('excipient2') else ''

    api_sm = get_sm(api_input)
    exc1_sm = get_sm(exc1_input)
    exc2_sm = get_sm(exc2_input) if exc2_input else None

    if not api_sm or not exc1_sm:
        return jsonify({'error': '无法解析分子结构'})

    a = _cmol(api_sm)
    e1 = _cmol(exc1_sm)
    e2 = _cmol(exc2_sm) if exc2_sm else None
    if not a or not e1:
        return jsonify({'error': '分子结构解析失败'})

    # Reaction SMARTS descriptions
    REACT_SHORT = {
        'est':'Ester','amd':'Amide','lac':'Lactam','lct':'Lactone',
        'act':'Acetal','imn':'Imine','ahd':'Anhydride',
        'pho':'Phenol','ani':'Aniline','thl':'Thiol','the':'Thioether',
        'ald':'Aldehyde','pal':'1° Alcohol','ben':'Benzylic','all':'Allyl',
        'nit':'Nitro','nox':'N-oxide','qui':'Quinone',
        'acl':'Ar-Cl','abr':'Ar-Br','aid':'Ar-I',
        'bdk':'β-diketone','cat':'Catechol','cax':'Carboxylic acid','pam':'1° Amine',
    }
    HYDROLYSIS = ['est','amd','lac','lct','act','imn','ahd']
    OXIDATION = ['pho','ani','thl','the','ald','pal','ben','all']
    PHOTOLYSIS = ['nit','nox','qui','acl','abr','aid']
    CHELATION = ['bdk','cat','cax','pam']
    PW_DESCS = [
        ('Maillard (amine+ester)', 'API amine reacts with excipient ester -> Schiff base/glycation. Common with lactose and other reducing sugars at high temp.'),
        ('Maillard (amine+aldehyde)', 'API amine reacts with excipient aldehyde -> browning/degradation. Highly reactive combination.'),
        ('Esterification (acid+alcohol)', 'API acid + excipient alcohol -> ester formation under heat. Consumes API.'),
        ('Transesterification (ester+alcohol)', 'API ester + excipient alcohol -> ester exchange. Common with polyol excipients.'),
        ('Amidation (ester+amine)', 'API ester + excipient amine -> amide formation. Bioconjugation reaction.'),
        ('Hydrolysis (ester+acid)', 'API ester + excipient acid -> acid-catalyzed ester cleavage. Greatly accelerated at low pH.'),
        ('Hydrolysis (amide+acid)', 'API amide + excipient acid -> acid-catalyzed amide hydrolysis.'),
        ('Hydrolysis (ester+base)', 'API ester + excipient base -> base-catalyzed saponification. Mg stearate common culprit.'),
        ('Hydrolysis (amide+base)', 'API amide + excipient base -> base-catalyzed amide hydrolysis.'),
        ('Oxidation (phenol+thioether)', 'API phenol + excipient thioether -> phenol oxidation catalyzed by peroxide from thioether.'),
        ('Oxidation (thiol+thioether)', 'API thiol + excipient thioether -> thiol oxidation & disulfide formation.'),
        ('Oxidation (aniline+thioether)', 'API aniline + excipient thioether -> quinone-imine formation.'),
        ('Oxidation (aldehyde+thioether)', 'API aldehyde + excipient thioether -> aldehyde oxidation to acid.'),
        ('Oxidation (benzylic+thioether)', 'API benzylic CH₂ + excipient thioether -> radical-mediated oxidation.'),
        ('Chelation (diketone+acid)', 'API β-diketone + excipient acid -> metal chelate complex. Can trap metal ions.'),
        ('Chelation (catechol+acid)', 'API catechol + excipient acid -> bidentate metal chelation + discoloration.'),
        ('Schiff base (aldehyde+amine)', 'API aldehyde + excipient amine -> imine formation. Reversible but colored.'),
        ('Acid-base (acid+base)', 'API acid + excipient base -> salt formation. May alter solubility and stability.'),
        # Expanded
        ('Nitrile hydrolysis (acid)', 'API nitrile + excipient acid -> nitrile hydrolysis to amide or carboxylic acid.'),
        ('Nitrile hydrolysis (base)', 'API nitrile + excipient base -> base-catalyzed nitrile hydrolysis.'),
        ('Epoxide opening (amine)', 'API epoxide + excipient amine -> β-amino alcohol via ring opening.'),
        ('Epoxide opening (alcohol)', 'API epoxide + excipient alcohol -> β-hydroxy ether via ring opening.'),
        ('Epoxide opening (thiol)', 'API epoxide + excipient thiol -> β-hydroxy thioether via ring opening.'),
        ('Epoxide opening (acid)', 'API epoxide + excipient acid -> acid-catalyzed epoxide ring opening.'),
        ('Michael addition (amine)', 'API α,β-unsat carbonyl + excipient amine -> 1,4-addition. Common with acrylates/maleimides.'),
        ('Michael addition (thiol)', 'API α,β-unsat carbonyl + excipient thiol -> 1,4-addition. Fast reaction with maleimides.'),
        ('Disulfide formation', 'API thiol + excipient thiol -> disulfide exchange or oxidation. Reversible but labile.'),
        ('Hemiacetal formation', 'API aldehyde + excipient alcohol -> hemiacetal. Reversible equilibrium.'),
        ('Acetal formation', 'API aldehyde + excipient diol -> cyclic acetal. pH-sensitive.'),
        ('Nitrosamine formation', 'API amine + excipient nitrite/nitrate -> N-nitrosamine. Toxicological concern.'),
        ('Alkylation (amine)', 'API alkyl halide + excipient amine -> N-alkylation. Consumes API.'),
        ('Alkylation (thiol)', 'API alkyl halide + excipient thiol -> S-alkylation.'),
        ('Quinone-amine (Michael)', 'API quinone + excipient amine -> 1,4-addition, colored adduct.'),
        ('β-lactam opening (amine)', 'API β-lactam + excipient amine -> β-amino amide ring opening.'),
        ('β-lactam opening (alcohol)', 'API β-lactam + excipient alcohol -> β-hydroxy amide ring opening.'),
        ('β-lactam hydrolysis', 'API β-lactam + aq. excipient -> β-amino acid. Inherent instability.'),
        ('Sulfonate ester formation', 'API sulfonic acid + excipient alcohol -> sulfonate ester. Potential alkylating agent.'),
        ('Phosphate ester formation', 'API phosphate + excipient alcohol -> phosphate ester. Reversible.'),
        ('Urea formation', 'API isocyanate + excipient amine -> substituted urea. Highly reactive.'),
        ('Carbamate formation', 'API isocyanate + excipient alcohol -> carbamate. Moisture-sensitive.'),
        ('Ether cleavage (acid)', 'API ether + excipient acid -> ether bond cleavage via protonation.'),
        ('β-elimination', 'API leaving group + excipient base -> elimination, forms double bond.'),
        ('Halogen hydrolysis', 'API alkyl halide + excipient base/water -> alcohol formation.'),
        ('Diazo coupling (phenol)', 'API diazonium + excipient phenol -> azo dye. Colored impurity formation.'),
        ('Diazo coupling (aniline)', 'API diazonium + excipient aniline -> azo dye. Colored impurity formation.'),
        ('Amide formation (acid+amine)', 'API acid + excipient amine -> amide condensation. Requires activation.'),
        ('Thioester formation', 'API acid + excipient thiol -> thioester. Labile under hydrolysis.'),
        ('Hydrazone (carbonyl+hydrazine)', 'API carbonyl + excipient hydrazine -> hydrazone. Stable conjugate.'),
        ('Hydrazone (aldehyde+hydrazine)', 'API aldehyde + excipient hydrazine -> hydrazone. Rapid condensation.'),
        ('Oxime (carbonyl+hydroxylamine)', 'API carbonyl + excipient hydroxylamine -> oxime. Common analytical derivatization.'),
        ('Oxime (aldehyde+hydroxylamine)', 'API aldehyde + excipient hydroxylamine -> oxime. Fast reaction.'),
        ('Aminal formation', 'API carbonyl + excipient amine -> gem-diamine. Reversible equilibrium.'),
        ('Enamine formation', 'API carbonyl + excipient secondary amine -> enamine. Water eliminated.'),
        ('Transamidation', 'API amide + excipient amine -> amide exchange. Slow under normal conditions.'),
        ('Anhydride aminolysis', 'API anhydride + excipient amine -> amide + acid. Highly reactive.'),
        ('Anhydride alcoholysis', 'API anhydride + excipient alcohol -> ester + acid.'),
        ('Anhydride hydrolysis', 'API anhydride + water -> 2× carboxylic acid. Moisture-sensitive.'),
        ('Carbamate hydrolysis', 'API carbamate + water/base -> amine + CO₂ + alcohol.'),
        ('Urea hydrolysis', 'API urea + water/base -> amine + CO₂ + amine.'),
        ('Thiourea formation', 'API isothiocyanate + excipient amine -> thiourea derivative.'),
        ('Isothiocyanate+alcohol', 'API isothiocyanate + excipient alcohol -> thiocarbamate.'),
        ('Azide reduction (thiol)', 'API azide + excipient thiol -> amine. Click-to-release.'),
        ('N-oxide reduction', 'API N-oxide + excipient thiol -> tertiary amine. Reductive excipient.'),
        ('Hofmann elimination', 'API quaternary ammonium + base -> alkene + tertiary amine.'),
        ('Epoxidation (alkene->epoxide)', 'API alkene + excipient peroxide -> epoxide. Genotoxic risk.'),
        ('Alcohol oxidation', 'API alcohol + excipient oxidizing agent -> aldehyde/acid.'),
        ('Ether peroxide (PEG)', 'API/Excipient ether + O₂ -> hydroperoxide. PEG degradation.'),
        ('Thiazolidine formation', 'API aldehyde + excipient cysteine/thiol-amine -> thiazolidine ring.'),
        ('Hemithioacetal', 'API aldehyde + excipient thiol -> hemithioacetal. Reversible.'),
        ('Disulfide-thiol exchange', 'API disulfide + excipient thiol -> mixed disulfide.'),
        ('Ketal formation', 'API ketone + excipient diol -> ketal. Acid-labile.'),
        ('Amidoxime formation', 'API nitrile + excipient hydroxylamine -> amidoxime.'),
        ('Semicarbazone', 'API carbonyl + excipient semicarbazide -> semicarbazone.'),
        ('Dehydration (acid-catalyzed)', 'API alcohol + excipient acid -> alkene + H₂O.'),
        ('Decarboxylation', 'API β-keto acid + base/heat -> CO₂ + ketone.'),
        ('Sulfoxide reduction', 'API sulfoxide + excipient thiol -> sulfide.'),
        ('Iminium formation', 'API aldehyde + excipient tertiary amine -> iminium salt. Reactive intermediate.'),
        ('Mannich reaction', 'API enolizable carbonyl + excipient amine -> β-amino carbonyl.'),
        ('Pinner reaction', 'API nitrile + excipient alcohol -> imidate ester.'),
        ('Aldol condensation', 'API enolizable carbonyl + excipient aldehyde -> β-hydroxy carbonyl.'),
        ('Knoevenagel condensation', 'API active methylene + excipient aldehyde -> α,β-unsaturated.'),
        ('Sulfide->sulfoxide', 'API thioether + excipient peroxide -> sulfoxide. Common oxidation.'),
        ('Sulfoxide->sulfone', 'API sulfoxide + excess peroxide -> sulfone. Further oxidation.'),
        ('Nitro reduction (thiol)', 'API nitro + excipient thiol -> amine. Reducing impurities.'),
        ('Nitro reduction (sugar)', 'API nitro + excipient reducing sugar -> amine.'),
        ('Aldehyde reduction (sugar)', 'API aldehyde + excipient reducing sugar -> alcohol.'),
        ('Thiol dimerization', 'API thiol + O₂/peroxide -> disulfide. Oxidative coupling.'),
        ('Phenol oxidative dimer', 'API phenol + oxidant -> biphenol/biphenol ether. Browning.'),
        ('Formaldehyde + amine', 'Formaldehyde from excipient degradation + API amine -> imine.'),
        ('Formaldehyde + amide', 'Formaldehyde + API amide -> N-hydroxymethylamide.'),
        ('Formaldehyde + phenol', 'Formaldehyde + API phenol -> hydroxymethylphenol.'),
        ('Sucrose hydrolysis', 'Sucrose + water/acid -> glucose + fructose. Invert sugar.'),
        ('PEG autoxidation', 'PEG + O₂ -> aldehydes, acids, esters. Chain degradation.'),
        ('Ascorbate oxidation', 'Vitamin C + O₂ -> dehydroascorbic acid. Antioxidant depletion.'),
        ('Polysorbate hydrolysis', 'Polysorbate ester + water -> fatty acid + sorbitan. Surfactant degradation.'),
        ('Lactone hydrolysis', 'API lactone + water -> hydroxy acid. Ring opening.'),
        ('Lactam hydrolysis', 'API cyclic amide + water -> amino acid. Ring opening.'),
        ('Sulfonamide hydrolysis', 'API sulfonamide + acid -> amine + sulfonic acid.'),
        ('Phosphonate hydrolysis', 'API phosphonate + water -> phosphoric acid.'),
        ('Boronic ester hydrolysis', 'API boronate ester + water -> boronic acid + alcohol.'),
        ('Barbiturate hydrolysis', 'API barbiturate + base -> ring-opened product.'),
        ('Amidine formation', 'API nitrile + excipient amine -> amidine.'),
        ('Phosphoramidate hydrolysis', 'API phosphoramidate + water -> P-N cleavage.'),
        ('Acyl glucuronide hydrolysis', 'API acyl glucuronide + water -> aglycone + glucuronic acid.'),
        ('Hypochlorite + phenol', 'Hypochlorite (disinfectant residue) + API phenol -> chlorophenol.'),
        ('Photo-Fries rearrangement', 'Aryl ester + UV light -> o/p-hydroxy ketone. Photolabile.'),
        ('Norrish type II cyclization', 'Ketone with γ-H + UV -> cyclobutanol. Photo-cyclization.'),
        ('Photodehalogenation', 'Aryl halide + UV light -> dehalogenation. Photolytic cleavage.'),
        ('Photo-isomerization', 'Alkene + UV light -> cis/trans isomerization. Photosensitivity.'),
        ('Photo-oxidation', 'Alkene + O₂ + light -> hydroperoxide. Photo-initiated autoxidation.'),
        ('EDTA metal chelation', 'EDTA + Ca/Mg/Fe/Zn -> strong chelate. Sequesters catalytic metals.'),
        ('Citrate metal chelation', 'Citric acid + di/trivalent metal -> soluble complex.'),
        ('Thiol-metal complex', 'Thiol + Hg/Ag/Cu/Fe -> stable mercaptide complex.'),
        ('Amine-metal complex', 'Amine + Cu/Fe/Zn -> coordination complex. May catalyze degradation.'),
        ('Acid-metal complex', 'Carboxylic acid + metal -> salt/complex. May alter solubility.'),
        ('Charge-transfer complex', 'Aromatic donor + acceptor -> colored complex.'),
        ('Lipid peroxidation', 'Unsaturated lipid + O₂ -> aldehyde. Rancidity.'),
        ('Strecker degradation', 'Carbonyl + amino acid -> aldehyde + CO₂ + amine. Browning.'),
        ('Epimerization (base)', 'Chiral center + base -> epimerization. Loss of stereopurity.'),
        ('Racemization', 'Chiral center via enolization -> racemate. Activity loss.'),
        ('O-dealkylation (oxidative)', 'Alkyl ether + oxidant -> alcohol + aldehyde.'),
        ('Salt formation (tertiary amine+acid)', 'Tertiary amine + acid -> ammonium salt.'),
        ('Azo reduction', 'Azo bond + thiol -> hydrazine/amine. Color loss.'),
        ('Vinyl sulfone + amine', 'Vinyl sulfone + amine -> Michael adduct. Irreversible.'),
        ('Enamine hydrolysis', 'Enamine + water -> carbonyl + amine. Reversible.'),
        ('Imine hydrolysis', 'Imine + water -> carbonyl + amine. pH-dependent.'),
        ('Hydrazone hydrolysis', 'Hydrazone + water -> carbonyl + hydrazine. Reversible.'),
        ('Oxime hydrolysis', 'Oxime + water -> carbonyl + hydroxylamine. pH-dependent.'),
        ('Acetal hydrolysis', 'Acetal + water + acid -> carbonyl + alcohol. pH-sensitive.'),
        ('Epoxide hydrolysis', 'Epoxide + water -> diol. Acid/base catalyzed.'),
        ('Michael addition (alcohol)', 'α,β-unsaturated carbonyl + alcohol -> ether. Slow.'),
        ('Michael addition (water)', 'α,β-unsaturated carbonyl + water -> alcohol hydrate.'),
        ('Thioester hydrolysis', 'Thioester + water -> acid + thiol. Labile.'),
        ('Boronate ester (diol)', 'Boronic acid + diol -> cyclic boronate ester. pH-sensitive.'),
    ]

    def _group_analysis(cmol):
        rv = cmol['react']
        groups = {'hydrolysis': [], 'oxidation': [], 'photolysis': [], 'chelation': []}
        for i, code in enumerate(['est','amd','lac','lct','act','imn','ahd',
                                   'pho','ani','thl','the','ald','pal','ben','all',
                                   'nit','nox','qui','acl','abr','aid',
                                   'bdk','cat','cax','pam']):
            if rv[i] > 0:
                cat_name = 'hydrolysis' if code in HYDROLYSIS else \
                           'oxidation' if code in OXIDATION else \
                           'photolysis' if code in PHOTOLYSIS else 'chelation'
                groups.setdefault(cat_name, []).append(f'{REACT_SHORT[code]} ({int(rv[i])})')
        groups['total_rings'] = int(rv[25])
        groups['aromatic_rings'] = int(rv[26])
        groups['hbd_hba_sum'] = int(rv[27])
        groups['custom_groups'] = {}
        cg_names = {0:'sulfonate',1:'sulfate',2:'phosphate',3:'nitrile',4:'epoxide',
                    5:'carbamate',6:'urea',7:'sulfonamide',8:'boronic_acid',9:'tert_amine',
                    10:'quat_N',11:'ether',12:'carbonyl',13:'azide',14:'hydrazine',
                    15:'hydroxylamine',16:'F',17:'Cl',18:'Br',19:'I'}
        for i, cnt in enumerate(cmol['cg']):
            if cnt > 0 and i in cg_names:
                groups['custom_groups'][cg_names[i]] = int(cnt)
        return groups

    def _pairwise(mol_a, mol_e):
        pw_vec = _pairwise_vec(mol_a, mol_e) if mol_e else [False]*len(PW_DESCS)
        risks = []
        for val, (name, desc) in zip(pw_vec, PW_DESCS):
            risks.append({'name': name, 'active': bool(val), 'description': desc})
        return risks

    api_groups = _group_analysis(a)
    exc1_groups = _group_analysis(e1)
    pw_risks = _pairwise(a['mol'], e1['mol'])

    result = {
        'api_groups': api_groups,
        'exc1_groups': exc1_groups,
        'pairwise_risks': pw_risks,
        'error': None,
    }

    if e2 is not None:
        result['exc2_groups'] = _group_analysis(e2)
        result['exc2_pairwise_risks'] = _pairwise(a['mol'], e2['mol'])

    return jsonify(result)


@app.route('/api/parse', methods=['POST'])
def api_parse():
    """AI 解析自然语言输入, 返回结构化字段(分词结果). """
    data = request.json
    text = data.get('text', '').strip()
    if not text:
        return jsonify({'error': '请输入文本'})

    sys.path.insert(0, os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', 'src'))
    from chat import ai_parse_formulation, parse_formulation_input

    parsed = ai_parse_formulation(text)
    if not parsed.get('api_name') or not parsed.get('exc1_name'):
        parsed = parse_formulation_input(text)

    return jsonify({'parsed': parsed, 'error': None})


@app.route('/api/chat/stream', methods=['POST'])
def api_chat_stream():
    data = request.json
    question = data.get('question', '').strip()
    api_name = data.get('api', '').strip()
    exc1_name = data.get('excipient1', '').strip()
    exc2_name = data.get('excipient2', '').strip() if data.get('excipient2') else None
    days = data.get('days', None)
    cond = data.get('condition', None)
    history = data.get('history', None)
    api_key = data.get('api_key', None)
    api_url = data.get('api_url', None)
    model = data.get('model', None)
    skill = data.get('skill', 'auto')
    confirmed = data.get('confirmed', False)
    page_url = data.get('page_url', '')
    page_content = data.get('page_content', '')

    if not question:
        return jsonify({'error': 'Empty question'})

    # Inject page content into question context for AI reading
    if page_url and page_content:
        ctx = f"\n\n[当前浏览的网页: {page_url}]\n{page_content[:8000]}\n[/网页内容]"
        question = question + ctx

    sys.path.insert(0, os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', 'src'))
    from chat import chat_stream

    def generate():
        for chunk in chat_stream(question, api_name, exc1_name, days, cond, history, api_key=api_key, api_url=api_url, model=model, skill=skill, exc2_name=exc2_name if exc2_name else None, confirmed=confirmed):
            if isinstance(chunk, dict):
                t = chunk.get('type')
                if t == 'status':
                    yield f"data: {json.dumps({'s': chunk['text']}, ensure_ascii=False)}\n\n"
                elif t == 'content':
                    yield f"data: {json.dumps({'c': chunk['text']}, ensure_ascii=False)}\n\n"
                elif t == 'phase':
                    yield f"data: {json.dumps({'p': {'phase': chunk.get('phase'), 'status': chunk.get('status'), 'text': chunk.get('text')}}, ensure_ascii=False)}\n\n"
                elif t == 'sources':
                    yield f"data: {json.dumps({'src': chunk['sources']}, ensure_ascii=False)}\n\n"
                elif t == 'web_action':
                    yield f"data: {json.dumps({'w': chunk}, ensure_ascii=False)}\n\n"
                elif t == 'parsed':
                    yield f"data: {json.dumps({'parsed': chunk['parsed']}, ensure_ascii=False)}\n\n"
            else:
                yield f"data: {json.dumps({'c': chunk}, ensure_ascii=False)}\n\n"

    return Response(stream_with_context(generate()), mimetype='text/event-stream')

@app.route('/api/chat/export-docx', methods=['POST'])
def api_export_docx():
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from urllib.parse import quote
    import datetime as _dt
    import re as _re

    data = request.json
    title = data.get('title', '实验设计方案')
    content = data.get('content', '')
    api_name = data.get('api_name', '')
    exc1_name = data.get('exc1_name', '')
    exc2_name = data.get('exc2_name', '')
    if not content:
        return jsonify({'error': '内容为空'}), 400

    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.8)

    def _run_p(text, size=11, bold=False, color=RGBColor(0x1d, 0x1d, 0x1f),
               space_after=6, space_before=0, alignment=None, italic=False):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.space_before = Pt(space_before)
        if alignment is not None:
            p.alignment = alignment
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.italic = italic
        run.font.color.rgb = color
        return p

    def _clean_md_text(text):
        """Strip markdown symbols that _add_inline_md doesn't handle: images, code, italic, strikethrough."""
        text = _re.sub(r'!\[.*?\]\(.*?\)', '', text)       # ![alt](url)
        text = _re.sub(r'(`+)(.*?)\1', r'\2', text)         # `code` / ``code``
        text = _re.sub(r'~~(.*?)~~', r'\1', text)            # ~~strikethrough~~
        text = _re.sub(r'(?<!\*)\*(?!\*)(.*?)(?<!\*)\*(?!\*)', r'\1', text)  # *italic* not **bold**
        text = _re.sub(r'(?<!_)_(?!_)(.*?)(?<!_)_(?!_)', r'\1', text)  # _italic_ not __
        text = _re.sub(r'^>\s+', '', text, flags=_re.MULTILINE)  # blockquotes
        text = _re.sub(r'^[-*_]{3,}\s*$', '', text, flags=_re.MULTILINE)  # hr
        return text

    def _add_inline_md(p, text, size=11, base_color=RGBColor(0x1d, 0x1d, 0x1f)):
        """Parse inline markdown (**bold**, [links](url)) and add formatted text to paragraph.
        Bold is handled first as a separate pass, then links, then cleanup."""
        text = _clean_md_text(text)
        # 1st pass: handle **bold** — split on **...** patterns
        bold_parts = _re.split(r'(\*\*.*?\*\*)', text)
        for bp in bold_parts:
            if bp.startswith('**') and bp.endswith('**'):
                r = p.add_run(bp[2:-2])
                r.font.size = Pt(size)
                r.font.bold = True
                r.font.color.rgb = base_color
            else:
                # 2nd pass: handle [links](url) in non-bold text
                link_parts = _re.split(r'(\[.*?\]\(.*?\))', bp)
                for lp in link_parts:
                    m = _re.match(r'\[(.+?)\]\((.+?)\)', lp)
                    if m:
                        r = p.add_run(m.group(1))
                        r.font.size = Pt(size)
                        r.font.color.rgb = RGBColor(0x00, 0x71, 0xe3)
                        r.font.underline = True
                        try:
                            from docx.opc.constants import RELATIONSHIP_TYPE as RT
                            rel = p.part.relate_to(m.group(2), RT.HYPERLINK, is_external=True)
                            r.hyperlink = rel
                        except:
                            pass
                    else:
                        # Strip any stray ** markers that weren't parsed
                        cleaned = lp.replace('**', '')
                        r = p.add_run(cleaned)
                        r.font.size = Pt(size)
                        r.font.color.rgb = base_color

    def _heading(level, text):
        h = doc.add_heading(text, level=level)
        return h

    def _bullet(text, size=11, indent=0):
        p = doc.add_paragraph(style='List Bullet')
        p.clear()
        run = p.add_run(text)
        run.font.size = Pt(size)
        return p

    def _set_heading_font(heading):
        for run in heading.runs:
            try:
                run.font.name = 'Microsoft YaHei'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            except:
                pass

    # ===================== COVER PAGE =====================
    _run_p('', size=11, space_after=60)
    _run_p('原辅料相容性', size=20, bold=True, color=RGBColor(0x15, 0x65, 0xc0),
           space_after=2, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    _run_p('实验方案设计报告', size=20, bold=True, color=RGBColor(0x1d, 0x1d, 0x1f),
           space_after=6, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    subtitle = f'基于 ML 预测与文献检索的全量分析'
    _run_p(subtitle, size=11, color=RGBColor(0x6b, 0x6b, 0x70),
           space_after=30, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # Metadata block
    exc_str = exc1_name
    if exc2_name:
        exc_str += f' + {exc2_name}'
    meta_lines = [
        f'检索对象: {api_name} / {exc_str}',
        f'生成日期: {_dt.datetime.now().strftime("%Y-%m-%d %H:%M")}',
        f'报告类型: AI 辅助原辅料相容性实验方案设计',
        f'分析方法: FP-XGB 模型预测 + RAG 知识库 + 文献检索',
    ]
    for ml in meta_lines:
        _run_p(ml, size=11, color=RGBColor(0x6b, 0x6b, 0x70), space_after=2,
               alignment=WD_ALIGN_PARAGRAPH.CENTER)

    div_p = doc.add_paragraph()
    div_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    div_run = div_p.add_run('-' * 45)
    div_run.font.color.rgb = RGBColor(0xcc, 0xcc, 0xd0)
    div_run.font.size = Pt(9)

    # ===================== PARSE & BUILD =====================
    lines = content.split('\n')
    ref_items = []
    in_refs = False

    def _detect_heading(line):
        m = _re.match(r'^(#{1,3})\s+(.+)$', line)
        if m:
            level = len(m.group(1))
            return level, m.group(2).strip()
        return None

    i = 0
    in_conclusion = False
    conclusion_buf = []

    def _flush_conclusion(buf):
        if not buf:
            return
        # Use a single-cell shaded table as a conclusion box
        from docx.enum.table import WD_TABLE_ALIGNMENT
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = table.cell(0, 0)
        cell.text = ''
        shading = cell._element.get_or_add_tcPr()
        shd = shading.makeelement(qn('w:shd'), {
            qn('w:fill'): 'e8f0fe',
            qn('w:val'): 'clear',
        })
        shading.append(shd)
        p0 = cell.paragraphs[0]
        r0 = p0.add_run('[DATA] 结论摘要')
        r0.font.size = Pt(11)
        r0.font.bold = True
        r0.font.color.rgb = RGBColor(0x15, 0x65, 0xc0)
        for b_line in buf:
            p = cell.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            _add_inline_md(p, b_line, size=11)

    while i < len(lines):
        line = lines[i]
        trimmed = line.strip()

        if not trimmed:
            i += 1
            continue

        # References section - capture all lines under "参考来源" / "References"
        if _re.search(r'^#{0,3}\s*[0-9.]*\s*参考来源|^#{0,3}\s*References|^\s*[0-9.]+\s*参考来源', trimmed):
            in_refs = True
            i += 1
            continue

        if in_refs:
            if _detect_heading(trimmed) or trimmed.startswith('—'):
                in_refs = False
                continue
            if trimmed.strip():
                ref_items.append(_clean_md_text(trimmed))
            i += 1
            continue

        # Conclusion summary special handling
        if '结论摘要' in trimmed:
            in_conclusion = True
            conclusion_buf = []
            i += 1
            continue
        # End conclusion on next heading
        if in_conclusion and _detect_heading(trimmed):
            _flush_conclusion(conclusion_buf)
            in_conclusion = False
            h = _detect_heading(trimmed)
            h_obj = _heading(h[0], h[1])
            _set_heading_font(h_obj)
            i += 1
            continue
        if in_conclusion:
            conclusion_buf.append(_clean_md_text(trimmed))
            i += 1
            continue

        h = _detect_heading(trimmed)
        if h:
            level, htext = h
            htext = _clean_md_text(htext)
            h_obj = _heading(level, htext)
            _set_heading_font(h_obj)
            i += 1
            continue

        # Markdown table: | ... | ... |
        if trimmed.startswith('|') and trimmed.endswith('|') and '|' in trimmed[1:-1]:
            rows = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                rows.append(lines[i].strip())
                i += 1
            if len(rows) >= 2 and _re.match(r'\|[\s\-:]+\|', rows[1]):
                headers = [h.strip() for h in rows[0].split('|')[1:-1]]
                data_rows = rows[2:]
                table = doc.add_table(rows=1 + len(data_rows), cols=len(headers))
                table.style = 'Light Grid Accent 1'
                # header row
                for ci, h in enumerate(headers):
                    cell = table.cell(0, ci)
                    cell.text = ''
                    run = cell.paragraphs[0].add_run(h)
                    run.font.size = Pt(10)
                    run.font.bold = True
                # data rows
                for ri, row_line in enumerate(data_rows):
                    cells = [c.strip() for c in row_line.split('|')[1:-1]]
                    for ci, cv in enumerate(cells):
                        if ci >= len(headers):
                            break
                        cell = table.cell(ri + 1, ci)
                        cell.text = ''
                        _add_inline_md(cell.paragraphs[0], cv, size=10)
            continue

        # Bullet / Numbered / Plain
        if trimmed.startswith('- ') or trimmed.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            p.clear()
            _add_inline_md(p, _re.sub(r'^[-*]\s', '', trimmed))
        elif _re.match(r'^\d+[\.\)]\s', trimmed):
            p = doc.add_paragraph(style='List Bullet')
            p.clear()
            _add_inline_md(p, _re.sub(r'^\d+[\.\)]\s', '', trimmed))
        elif _re.match(r'^\(?\d+\)\s', trimmed):
            p = doc.add_paragraph(style='List Bullet')
            p.clear()
            _add_inline_md(p, _re.sub(r'^\(?\d+\)\s', '', trimmed))
        else:
            p = doc.add_paragraph()
            _add_inline_md(p, trimmed)

        i += 1

    # ===================== REFERENCES =====================
    if ref_items:
        doc.add_page_break()
        _heading(1, '4. 参考来源')
        for ref in ref_items:
            clean = ref.lstrip('- *')

            # Parse numbered items: "1. text [title](url)"
            m = _re.match(r'^(\d+[\.\)\s]*)', clean)
            if m:
                clean = clean[m.end():]

            p = doc.add_paragraph(style='List Bullet')
            p.clear()
            _add_inline_md(p, clean.strip(), size=11)

    # ===================== FOOTER =====================
    from docx.shared import Pt as _Pt
    doc.add_paragraph()
    p_footer = doc.add_paragraph()
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_footer = p_footer.add_run('— 本报告由 AI 辅助生成, 仅供参考, 请以实验数据为准 —')
    r_footer.font.size = Pt(9)
    r_footer.font.color.rgb = RGBColor(0xaa, 0xaa, 0xaf)
    r_footer.font.italic = True

    buf = __import__('io').BytesIO()
    doc.save(buf)
    buf.seek(0)
    safe_name = quote(f'{title}.docx')
    return Response(buf.read(), mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                    headers={'Content-Disposition': f'attachment; filename*=UTF-8\'\'{safe_name}'})

@app.route('/api/proxy-page')
def api_proxy_page():
    url = request.args.get('url', '').strip()
    if not url:
        return 'Missing url', 400
    try:
        r = requests.get(url, timeout=15, headers={
            'User-Agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/134.0.0.0 Safari/537.36',
            'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8',
            'Accept-Language': 'en-US,en;q=0.5',
        })
        resp = Response(r.content, status=r.status_code, content_type=r.headers.get('Content-Type', 'text/html'))
        # Strip X-Frame-Options and CSP to allow iframe embedding
        for h in ('X-Frame-Options', 'Content-Security-Policy', 'X-Content-Type-Options'):
            if h in resp.headers:
                del resp.headers[h]
        # Handle non-200 status codes
        if r.status_code >= 400:
            return f'<html><body style="font-family:sans-serif;padding:20px"><h2 style="color:#c62828">无法访问此页面</h2><p>HTTP {r.status_code}</p><p><a href="{url}" target="_blank" style="color:#0071e3">在新标签页中打开 ↗</a></p></body></html>', 200, {'Content-Type': 'text/html'}
        # Inject <base> tag + click interceptor
        if 'text/html' in (r.headers.get('Content-Type', '')):
            html = r.text
            base = f'<base href="{url.rstrip("/")}/">'
            interceptor = (
                '<script>'
                'document.addEventListener("click",function(e){'
                'var t=e.target.closest("a");'
                'if(t&&t.href&&!t.hasAttribute("data-wv-skip")){'
                'e.preventDefault();'
                'parent.postMessage({type:"wv-navigate",url:t.href},"*")'
                '}})</script>'
            )
            if '</head>' in html:
                html = html.replace('</head>', base + interceptor + '</head>')
            else:
                html = base + interceptor + html
            resp.set_data(html)
        return resp
    except Exception as e:
        return f'<html><body style="font-family:sans-serif;padding:20px"><h2 style="color:#c62828">无法加载页面</h2><p>{e}</p><p>请尝试<a href="{url}" target="_blank" style="color:#0071e3">在新标签页中打开</a></p></body></html>', 200, {'Content-Type': 'text/html'}

@app.route('/api/fetch-page', methods=['POST'])
def api_fetch_page():
    data = request.json
    url = (data.get('url') or '').strip()
    if not url:
        return jsonify({'error': 'Empty URL'}), 400
    try:
        import bs4
        r = requests.get(url, timeout=10,
                         headers={'User-Agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36'})
        if r.status_code != 200:
            return jsonify({'error': f'HTTP {r.status_code}', 'content': '', 'word_count': 0})
        soup = bs4.BeautifulSoup(r.text, 'html.parser')
        for tag in soup(['script', 'style', 'nav', 'footer', 'header', 'aside']):
            tag.decompose()
        text = soup.get_text(separator='\n', strip=True)
        text = '\n'.join(line.strip() for line in text.split('\n') if line.strip())
        title = soup.title.string.strip() if soup.title and soup.title.string else ''
        return jsonify({
            'url': url,
            'title': title,
            'content': text[:15000],
            'word_count': len(text[:15000]),
        })
    except Exception as e:
        return jsonify({'error': str(e), 'content': '', 'word_count': 0})

# -- Feedback API --
@app.route('/api/feedback', methods=['POST'])
def api_feedback():
    data = request.json or {}
    content = (data.get('content') or '').strip()
    if not content:
        return jsonify({'ok': False, 'error': '请填写反馈内容'})
    feedback_file = os.path.join(_DATA_DIR, 'feedbacks.json')
    feedbacks = []
    if os.path.exists(feedback_file):
        try:
            with open(feedback_file, 'r') as f:
                feedbacks = json.load(f)
        except:
            feedbacks = []
    feedbacks.append({
        'id': len(feedbacks) + 1,
        'content': content,
        'time': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
    })
    with open(feedback_file, 'w') as f:
        json.dump(feedbacks, f, ensure_ascii=False, indent=2)
    return jsonify({'ok': True})


# -- Chat History API --
@app.route('/api/chat/save-history', methods=['POST'])
def api_save_history():
    data = request.json or {}
    messages = data.get('messages', [])
    skill = data.get('skill', 'auto')
    if not messages:
        return jsonify({'ok': False, 'error': 'No messages'})
    history_dir = os.path.join(_DATA_DIR, 'histories')
    os.makedirs(history_dir, exist_ok=True)
    record = {
        'id': datetime.now().strftime('%Y%m%d_%H%M%S'),
        'user': user,
        'skill': skill,
        'time': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
        'messages': messages,
    }
    safe_user = user.replace('/', '_')
    fpath = os.path.join(history_dir, f'{safe_user}_{record["id"]}.json')
    with open(fpath, 'w') as f:
        json.dump(record, f, ensure_ascii=False, indent=2)
    return jsonify({'ok': True})


@app.route('/api/admin/histories')
def api_admin_histories():
    history_dir = os.path.join(_DATA_DIR, 'histories')
    if not os.path.exists(history_dir):
        return jsonify({'histories': []})
    user_records = []
    for fn in sorted(os.listdir(history_dir), reverse=True):
        if fn.endswith('.json') and fn.startswith(f'{current_user}_'):
            try:
                with open(os.path.join(history_dir, fn), 'r') as f:
                    record = json.load(f)
                    record['preview'] = ''
                    for m in record.get('messages', []):
                        if m.get('role') == 'user':
                            record['preview'] = m.get('content', '')[:80]
                            break
                    user_records.append(record)
            except:
                pass
    return jsonify({'histories': user_records})


@app.route('/admin')
def admin_page():
    return '''
<!DOCTYPE html><html lang="zh-CN"><head>
<meta charset="utf-8"><meta name="viewport" content="width=device-width, initial-scale=1">
<title>后台管理 - 聊天记录</title>
<style>
* { margin:0; padding:0; box-sizing:border-box; }
body { font-family:-apple-system,BlinkMacSystemFont,sans-serif; background:#f5f5f7; color:#1d1d1f; padding:20px; }
h1 { font-size:22px; margin-bottom:20px; }
.record { background:#fff; border-radius:12px; padding:16px; margin-bottom:12px; box-shadow:0 2px 8px rgba(0,0,0,0.04); }
.record .meta { font-size:12px; color:#86868b; margin-bottom:6px; }
.record .preview { font-size:14px; margin-bottom:8px; }
.record .msgs { display:none; font-size:13px; line-height:1.6; border-top:1px solid #eee; padding-top:10px; margin-top:10px; }
.record .msgs.open { display:block; }
.record .msg { margin:4px 0; padding:6px 10px; border-radius:8px; }
.record .msg.user { background:#e3f2fd; }
.record .msg.assistant { background:#f5f5f5; }
.record .toggle { color:#0071e3; cursor:pointer; font-size:13px; }
</style></head><body>
<h1>[LIST] 我的聊天记录</h1>
<div id="root">Loading...</div>
<script>
fetch('/api/admin/histories').then(r=>r.json()).then(d=>{
  if(!d.histories||!d.histories.length) return document.getElementById('root').innerHTML='<p style="color:#86868b">暂无记录</p>';
  let html='';
  for(const h of d.histories){
    const msgs=JSON.stringify(h.messages||[]).replace(/</g,'&lt;');
    html+=`<div class="record">
      <div class="meta">👤 ${h.user} | [TARGET] ${h.skill||'auto'} | 🕐 ${h.time}</div>
      <div class="preview">${h.preview||'(空)'}</div>
      <span class="toggle" onclick="this.nextElementSibling.classList.toggle('open');this.textContent=this.nextElementSibling.classList.contains('open')?'收起':'展开'">展开</span>
      <div class="msgs">`;
    for(const m of (h.messages||[])){
      const role=m.role||'';
      const content=(m.content||'').replace(/</g,'&lt;').replace(/\n/g,'<br>');
      html+=`<div class="msg ${role}"><strong>${role==='user'?'👤 用户':'🤖 AI'}:</strong><br>${content}</div>`;
    }
    html+=`</div></div>`;
  }
  document.getElementById('root').innerHTML=html;
}).catch(()=>document.getElementById('root').innerHTML='<p style="color:#c62828">加载失败</p>');
</script></body></html>
'''


def warmup():
    _PRJ = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    def _load_everything():
        import time; t0 = time.time()
        # 1. Load ML models (XGBoost) so first request is fast
        print("[BG] Loading ML models...")
        try:
            from predictor import load as pred_load
            pred_load()
            print(f"  ML models OK ({time.time()-t0:.0f}s)")
        except Exception as e:
            print(f"  ML models skip: {e}")

        # 2. Load Knowledge Base
        t1 = time.time()
        print("[BG] Loading Knowledge Base...")
        try:
            sys.path.insert(0, os.path.join(_PRJ, 'src'))
            from knowledge_base import get_kb
            get_kb()
            print(f"  Knowledge Base OK ({time.time()-t1:.0f}s)")
        except Exception as e:
            print(f"  Knowledge Base skip: {e}")
        print(f"[BG] Warmup done ({time.time()-t0:.0f}s total)")
    import threading
    threading.Thread(target=_load_everything, daemon=True).start()

if __name__ == '__main__':
    warmup()
    print("\n" + "="*50)
    print("  药物-辅料相容性预测系统")
    print("  访问: http://localhost:8080")
    print("="*50 + "\n")
    app.run(debug=True, host='0.0.0.0', port=8080)
